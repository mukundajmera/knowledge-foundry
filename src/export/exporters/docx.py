"""Knowledge Foundry — DOCX Exporter.

Exports entities to Microsoft Word format (.docx).
Uses python-docx for document generation.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import TYPE_CHECKING

from src.export.base import BaseExporter, ExportableEntity
from src.export.models import (
    EntityType,
    ExportableConversation,
    ExportableEvaluationReport,
    ExportableMessage,
    ExportableRAGRun,
    ExportContext,
    ExportOptions,
    ExportResult,
)

if TYPE_CHECKING:
    pass

logger = logging.getLogger(__name__)


class DOCXExporter(BaseExporter):
    """Export entities to Microsoft Word format.

    Creates well-formatted DOCX documents with proper headings,
    paragraphs, and tables.
    """

    def __init__(self) -> None:
        """Initialize DOCX exporter and check for python-docx."""
        self._docx_available = self._check_docx()

    def _check_docx(self) -> bool:
        """Check if python-docx is available."""
        try:
            import docx  # noqa: F401
            return True
        except ImportError:
            logger.warning(
                "python-docx not installed. DOCX export will not be available. "
                "Install with: pip install python-docx"
            )
            return False

    @property
    def format_id(self) -> str:
        return "docx"

    @property
    def label(self) -> str:
        return "Word Document"

    @property
    def description(self) -> str:
        if self._docx_available:
            return "Export as a Microsoft Word document (.docx)"
        return "Export as a Word document (requires python-docx)"

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def extension(self) -> str:
        return ".docx"

    @property
    def supported_entity_types(self) -> list[EntityType]:
        return [
            EntityType.CONVERSATION,
            EntityType.MESSAGE,
            EntityType.RAG_RUN,
            EntityType.EVALUATION_REPORT,
        ]

    def generate(
        self,
        entity: ExportableEntity,
        options: ExportOptions,
        context: ExportContext,
    ) -> ExportResult:
        """Generate DOCX export."""
        if not self._docx_available:
            return ExportResult(
                success=False,
                error="DOCX generation requires python-docx. Install with: pip install python-docx",
            )

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Set up document styles
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            if isinstance(entity, ExportableConversation):
                self._export_conversation(doc, entity, options)
            elif isinstance(entity, ExportableMessage):
                self._export_message(doc, entity, options)
            elif isinstance(entity, ExportableRAGRun):
                self._export_rag_run(doc, entity, options)
            elif isinstance(entity, ExportableEvaluationReport):
                self._export_evaluation_report(doc, entity, options)
            else:
                return ExportResult(
                    success=False,
                    error=f"Unsupported entity type: {type(entity).__name__}",
                )

            # Add footer
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Exported from Knowledge Foundry on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            content = buffer.getvalue()

            if options.anonymize_user:
                # For DOCX, we need to reload and process
                content = self._anonymize_docx(content)

            return ExportResult(
                success=True,
                content=content,
                mime_type=self.mime_type,
            )
        except Exception as e:
            logger.exception("DOCX generation failed: %s", e)
            return ExportResult(
                success=False,
                error=f"DOCX generation failed: {str(e)}",
            )

    def _format_datetime(self, dt: datetime) -> str:
        """Format datetime for display."""
        return dt.strftime("%Y-%m-%d %H:%M:%S UTC")

    def _add_heading(self, doc: "Document", text: str, level: int = 1) -> None:
        """Add a heading to the document."""
        doc.add_heading(text, level=level)

    def _add_metadata_table(self, doc: "Document", items: list[tuple[str, str]]) -> None:
        """Add a metadata table."""
        from docx.shared import Pt

        table = doc.add_table(rows=len(items), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(items):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold the label
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True

    def _export_conversation(
        self,
        doc: "Document",
        conv: ExportableConversation,
        options: ExportOptions,
    ) -> None:
        """Export a conversation to DOCX."""
        from docx.shared import Pt, RGBColor

        self._add_heading(doc, conv.title, level=1)

        if options.include_metadata:
            self._add_heading(doc, "Conversation Details", level=2)
            meta_items = [
                ("ID", conv.id),
                ("Created", self._format_datetime(conv.created_at)),
                ("Updated", self._format_datetime(conv.updated_at)),
            ]
            if conv.model_info:
                meta_items.append(("Model", conv.model_info))
            self._add_metadata_table(doc, meta_items)
            doc.add_paragraph()

        self._add_heading(doc, "Messages", level=2)

        for msg in conv.messages:
            role_display = "User" if msg.role == "user" else "Assistant"
            if msg.role == "system":
                role_display = "System"

            # Add role as bold heading
            para = doc.add_paragraph()
            role_run = para.add_run(f"{role_display}: ")
            role_run.bold = True
            if msg.role == "user":
                role_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue

            if options.include_metadata:
                meta_run = para.add_run(f" [{self._format_datetime(msg.timestamp)}]")
                meta_run.font.size = Pt(9)
                meta_run.font.color.rgb = RGBColor(107, 114, 128)  # Gray

            # Add content
            doc.add_paragraph(msg.content)

        # Citations
        if options.include_citations:
            all_citations = []
            for msg in conv.messages:
                all_citations.extend(msg.citations)

            if all_citations:
                self._add_heading(doc, "Sources", level=2)
                seen_ids = set()
                for i, cite in enumerate(all_citations, 1):
                    if cite.document_id in seen_ids:
                        continue
                    seen_ids.add(cite.document_id)
                    para = doc.add_paragraph()
                    para.add_run(f"{i}. ").bold = True
                    para.add_run(cite.title).bold = True
                    if cite.section:
                        doc.add_paragraph(f"   Section: {cite.section}")
                    doc.add_paragraph(f"   Relevance: {cite.relevance_score:.0%}")

    def _export_message(
        self,
        doc: "Document",
        msg: ExportableMessage,
        options: ExportOptions,
    ) -> None:
        """Export a single message to DOCX."""
        role_display = "User Message" if msg.role == "user" else "Assistant Response"
        if msg.role == "system":
            role_display = "System Message"

        self._add_heading(doc, role_display, level=1)

        if options.include_metadata:
            self._add_heading(doc, "Details", level=2)
            meta_items = [
                ("ID", msg.id),
                ("Timestamp", self._format_datetime(msg.timestamp)),
            ]
            if msg.model:
                meta_items.append(("Model", msg.model))
            if msg.confidence is not None:
                meta_items.append(("Confidence", f"{msg.confidence:.0%}"))
            if msg.latency_ms:
                meta_items.append(("Latency", f"{msg.latency_ms}ms"))
            if msg.cost_usd:
                meta_items.append(("Cost", f"${msg.cost_usd:.4f}"))
            self._add_metadata_table(doc, meta_items)
            doc.add_paragraph()

        self._add_heading(doc, "Content", level=2)
        doc.add_paragraph(msg.content)

        if options.include_citations and msg.citations:
            self._add_heading(doc, "Citations", level=2)
            for i, cite in enumerate(msg.citations, 1):
                para = doc.add_paragraph()
                para.add_run(f"{i}. ").bold = True
                para.add_run(cite.title).bold = True
                doc.add_paragraph(f"   Relevance: {cite.relevance_score:.0%}")

    def _export_rag_run(
        self,
        doc: "Document",
        run: ExportableRAGRun,
        options: ExportOptions,
    ) -> None:
        """Export a RAG run to DOCX."""
        self._add_heading(doc, "RAG Query Result", level=1)

        if options.include_metadata:
            self._add_heading(doc, "Run Details", level=2)
            meta_items = [
                ("ID", run.id),
                ("Timestamp", self._format_datetime(run.timestamp)),
            ]
            if run.model:
                meta_items.append(("Model", run.model))
            if run.model_tier:
                meta_items.append(("Tier", run.model_tier))
            meta_items.append(("Latency", f"{run.latency_ms}ms"))
            meta_items.append(("Tokens", f"{run.input_tokens} in / {run.output_tokens} out"))
            if run.cost_usd:
                meta_items.append(("Cost", f"${run.cost_usd:.4f}"))
            if run.confidence is not None:
                meta_items.append(("Confidence", f"{run.confidence:.0%}"))
            self._add_metadata_table(doc, meta_items)
            doc.add_paragraph()

        self._add_heading(doc, "Query", level=2)
        para = doc.add_paragraph()
        para.add_run(run.query).italic = True

        self._add_heading(doc, "Answer", level=2)
        doc.add_paragraph(run.answer)

        if run.contexts:
            self._add_heading(doc, "Retrieved Context", level=2)
            for i, ctx in enumerate(run.contexts, 1):
                doc.add_heading(f"Context {i}: {ctx.title}", level=3)
                doc.add_paragraph(f"Score: {ctx.score:.0%}")
                text_preview = ctx.text[:500] + ("..." if len(ctx.text) > 500 else "")
                para = doc.add_paragraph()
                para.add_run(text_preview).italic = True

        if options.include_citations and run.citations:
            self._add_heading(doc, "Citations", level=2)
            for cite in run.citations:
                para = doc.add_paragraph()
                para.add_run(cite.title).bold = True
                para.add_run(f" (relevance: {cite.relevance_score:.0%})")

        if run.evaluation_metrics:
            self._add_heading(doc, "Evaluation Metrics", level=2)
            table = doc.add_table(rows=len(run.evaluation_metrics) + 1, cols=2)
            table.style = "Table Grid"
            # Header
            table.rows[0].cells[0].text = "Metric"
            table.rows[0].cells[1].text = "Value"
            for cell in table.rows[0].cells:
                for run_p in cell.paragraphs[0].runs:
                    run_p.bold = True
            # Data
            for i, (name, value) in enumerate(run.evaluation_metrics.items(), 1):
                table.rows[i].cells[0].text = name
                table.rows[i].cells[1].text = f"{value:.2f}"

    def _export_evaluation_report(
        self,
        doc: "Document",
        report: ExportableEvaluationReport,
        options: ExportOptions,
    ) -> None:
        """Export an evaluation report to DOCX."""
        from docx.shared import RGBColor

        self._add_heading(doc, report.title, level=1)

        if options.include_metadata:
            self._add_heading(doc, "Report Details", level=2)
            meta_items = [
                ("ID", report.id),
                ("Generated", self._format_datetime(report.timestamp)),
            ]
            if report.model_info:
                meta_items.append(("Model", report.model_info))
            meta_items.append(("Dataset Size", str(report.dataset_size)))
            meta_items.append(("Passed", str(report.passed_count)))
            meta_items.append(("Failed", str(report.failed_count)))
            if report.overall_score is not None:
                meta_items.append(("Overall Score", f"{report.overall_score:.2%}"))
            self._add_metadata_table(doc, meta_items)
            doc.add_paragraph()

        self._add_heading(doc, "Metrics", level=2)
        table = doc.add_table(rows=len(report.metrics) + 1, cols=4)
        table.style = "Table Grid"
        # Header
        headers = ["Metric", "Value", "Threshold", "Status"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        # Data
        for i, metric in enumerate(report.metrics, 1):
            table.rows[i].cells[0].text = metric.name
            table.rows[i].cells[1].text = f"{metric.value:.2f}"
            table.rows[i].cells[2].text = f"{metric.threshold:.2f}" if metric.threshold else "-"
            if metric.passed is True:
                table.rows[i].cells[3].text = "✓ Pass"
            elif metric.passed is False:
                table.rows[i].cells[3].text = "✗ Fail"
            else:
                table.rows[i].cells[3].text = "-"

        if report.examples:
            self._add_heading(doc, "Example Cases", level=2)
            for i, ex in enumerate(report.examples, 1):
                status = "✓ Pass" if ex.passed else "✗ Fail"
                doc.add_heading(f"Example {i} - {status}", level=3)
                para = doc.add_paragraph()
                para.add_run("Query: ").bold = True
                para.add_run(ex.query)
                if ex.expected:
                    para = doc.add_paragraph()
                    para.add_run("Expected: ").bold = True
                    para.add_run(ex.expected)
                para = doc.add_paragraph()
                para.add_run("Actual: ").bold = True
                para.add_run(ex.actual)
                if ex.notes:
                    para = doc.add_paragraph()
                    para.add_run("Notes: ").italic = True
                    para.add_run(ex.notes).italic = True

    def _anonymize_docx(self, content: bytes) -> bytes:
        """Anonymize content in a DOCX file."""
        import io
        from docx import Document

        buffer = io.BytesIO(content)
        doc = Document(buffer)

        # Process all paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = self._anonymize_content(run.text)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = self._anonymize_content(run.text)

        # Save back to bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
